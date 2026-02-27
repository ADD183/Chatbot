from celery import Celery
from sqlalchemy.orm import Session
from sqlalchemy import func
import os
from dotenv import load_dotenv
import fitz  # PyMuPDF
from typing import List, Tuple
import json
import re
from datetime import datetime
from docx import Document as DocxDocument
import traceback

from database import SessionLocal
from models import Document
from models import Document, EnqueueAudit
from gemini_service import gemini_service
import logging

logger = logging.getLogger(__name__)

load_dotenv()

REDIS_URL = os.getenv("REDIS_URL", "redis://localhost:6379/0")
CHUNK_SIZE = int(os.getenv('CHUNK_SIZE', '500'))
CHUNK_OVERLAP = int(os.getenv('CHUNK_OVERLAP', '50'))

# Initialize Celery
celery_app = Celery(
    "chatbot_worker",
    broker=REDIS_URL,
    backend=REDIS_URL
)

celery_app.conf.update(
    task_serializer='json',
    accept_content=['json'],
    result_serializer='json',
    timezone='UTC',
    enable_utc=True,
    task_track_started=True,
    task_time_limit=3600,  # 1 hour max per task
    worker_prefetch_multiplier=1,
)

# Helper to append debug entries to a host-visible file in uploads/
def _append_debug_log(line: str):
    try:
        now = datetime.utcnow().isoformat()
        uploads_path = os.path.join(os.path.dirname(__file__), 'uploads')
        os.makedirs(uploads_path, exist_ok=True)
        with open(os.path.join(uploads_path, 'worker_debug.log'), 'a', encoding='utf-8') as fh:
            fh.write(f"{now} - {line}\n")
    except Exception:
        pass

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[Tuple[str, int, int]]:
    """
    Split text into overlapping chunks and return chunk offsets.

    Returns a list of tuples: (chunk_text, start_char, end_char)
    """
    if not text or not text.strip():
        return []

    text = text.replace('\x00', ' ')

    # Clean the text but preserve original indexing by collapsing whitespace
    # into single spaces so offsets remain meaningful for the cleaned text.
    text = re.sub(r'\s+', ' ', text).strip()

    chunks: List[Tuple[str, int, int]] = []
    start = 0
    text_length = len(text)

    while start < text_length:
        end = start + chunk_size

        # If this is not the last chunk, try to break at a sentence or word boundary
        if end < text_length:
            sentence_end = max(
                text.rfind('. ', start, end),
                text.rfind('! ', start, end),
                text.rfind('? ', start, end)
            )

            if sentence_end > start:
                end = sentence_end + 1
            else:
                space_pos = text.rfind(' ', start, end)
                if space_pos > start:
                    end = space_pos

        # Defensive guard: ensure end is always ahead of start to avoid infinite loop
        if end <= start:
            end = start + max(1, chunk_size)

        chunk = text[start:end].strip()
        if chunk:
            chunks.append((chunk, start, end))

        # Move start position with overlap, but guarantee forward progress
        if end < text_length:
            next_start = end - overlap
            # If overlap is >= chunk_size or next_start doesn't advance, force a minimal advance
            if next_start <= start:
                next_start = start + max(1, chunk_size - overlap)
            start = next_start
        else:
            start = text_length

    return chunks

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If PDF extraction fails
    """
    try:
        doc = fitz.open(file_path)
        text_parts = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                text_parts.append(text)
        
        full_text = "\n\n".join(text_parts).replace('\x00', ' ')
        return full_text
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from TXT file
    
    Args:
        file_path: Path to TXT file
        
    Returns:
        File content
        
    Raises:
        Exception: If file reading fails
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read().replace('\x00', ' ')
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read().replace('\x00', ' ')
        except Exception as e:
            raise Exception(f"Failed to read TXT file: {str(e)}")
    except Exception as e:
        raise Exception(f"Failed to read TXT file: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = DocxDocument(file_path)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()).replace('\x00', ' ')
        if not text.strip():
            raise Exception("No readable text found in DOCX")
        return text
    except Exception as e:
        raise Exception(f"Failed to read DOCX file: {str(e)}")

@celery_app.task(bind=True, max_retries=3)
def process_document(
    self,
    file_path: str,
    filename: str,
    file_type: str,
    client_id: int
) -> dict:
    """
    Background task to process uploaded documents
    
    Args:
        file_path: Path to uploaded file
        filename: Original filename
        file_type: Type of file ('pdf' or 'txt')
        client_id: Client ID for multi-tenancy
        
    Returns:
        Dictionary with processing results
    """
    db: Session = SessionLocal()
    
    try:
        try:
            req_id = getattr(self.request, 'id', None)
        except Exception:
            req_id = None
        logger.info("process_document START: celery_request_id=%s, filename=%s, file_path=%s, file_type=%s, client_id=%s", req_id, filename, file_path, file_type, client_id)
        try:
            _append_debug_log(f"START: req_id={req_id}, filename={filename}, file_path={file_path}, file_type={file_type}, client_id={client_id}")
        except Exception:
            pass
        # Update audit row (mark started)
        try:
            if req_id:
                audit = db.query(EnqueueAudit).filter(EnqueueAudit.task_id == req_id).first()
            else:
                audit = db.query(EnqueueAudit).filter(EnqueueAudit.file_path == file_path, EnqueueAudit.filename == filename).order_by(EnqueueAudit.created_at.desc()).first()
            if audit:
                audit.status = 'started'
                audit.started_at = func.now()
                db.add(audit)
                db.commit()
                _append_debug_log(f"AUDIT_STARTED: task_id={audit.task_id}, file={filename}, client_id={audit.client_id}")
        except Exception:
            try:
                db.rollback()
            except Exception:
                pass
        if file_type == 'pdf':
            doc = fitz.open(file_path)
            total_pages = len(doc)
            logger.debug("Processing PDF %s with %s pages", filename, total_pages)
            
            chunk_idx = 0
            processed_chunks = 0
            
            for page_num in range(total_pages):
                page = doc[page_num]
                page_text = page.get_text()

                if not page_text.strip():
                    continue

                # Clean and chunk page text
                page_text = re.sub(r'\s+', ' ', page_text).strip()
                page_chunks = chunk_text(page_text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP)

                # Batch embeddings for this page
                batch = []
                BATCH_SIZE = int(os.getenv('EMBED_BATCH_SIZE', '16'))

                for start_chunk in page_chunks:
                    chunk_text_val, start_char, end_char = start_chunk
                    batch.append((chunk_text_val, start_char, end_char))

                    if len(batch) >= BATCH_SIZE:
                        texts = [b[0] for b in batch]
                        embeddings = gemini_service.generate_embeddings(texts)

                        for emb_vec, (chunk, s_char, e_char) in zip(embeddings, batch):
                            try:
                                if isinstance(emb_vec, dict):
                                    emb_vec = emb_vec.get('embedding', emb_vec)
                                try:
                                    emb_vec = emb_vec.tolist()
                                except:
                                    pass
                                emb_vec = [float(x) for x in emb_vec]

                                metadata = {
                                    'page': page_num + 1,
                                    'start_char': int(s_char),
                                    'end_char': int(e_char)
                                }

                                document = Document(
                                    client_id=client_id,
                                    filename=filename,
                                    file_type=file_type,
                                    chunk_text=chunk,
                                    chunk_index=chunk_idx,
                                    embedding=emb_vec,
                                    doc_metadata=json.dumps(metadata)
                                )
                                try:
                                    # Debug: log the client_id and chunk index right before DB insert
                                    logger.debug("DEBUG INSERT: client_id=%s, filename=%s, chunk_index=%s", client_id, filename, chunk_idx)
                                    _append_debug_log(f"DEBUG_INSERT: client_id={client_id}, filename={filename}, chunk_index={chunk_idx}")
                                except Exception:
                                    pass
                                try:
                                    db.add(document)
                                    db.flush()
                                except Exception as e:
                                    logger.warning("VECTOR INSERT FAILED: %s", e)
                                chunk_idx += 1
                                processed_chunks += 1
                            except Exception as e:
                                logger.exception("Error processing chunk %s on page %s: %s", chunk_idx, page_num, e)

                        # Commit the batch
                        db.commit()
                        logger.debug("Committed %s chunks for %s (page %s/%s)", processed_chunks, filename, page_num+1, total_pages)
                        try:
                            _append_debug_log(f"DB_COMMIT: filename={filename}, page={page_num+1}, committed_chunks={processed_chunks}")
                        except Exception:
                            pass
                        batch = []

                # Flush remaining items in batch
                if batch:
                    texts = [b[0] for b in batch]
                    embeddings = gemini_service.generate_embeddings(texts)
                    for emb_vec, (chunk, s_char, e_char) in zip(embeddings, batch):
                        try:
                            if isinstance(emb_vec, dict):
                                emb_vec = emb_vec.get('embedding', emb_vec)
                            try:
                                emb_vec = emb_vec.tolist()
                            except:
                                pass
                            emb_vec = [float(x) for x in emb_vec]

                            metadata = {
                                'page': page_num + 1,
                                'start_char': int(s_char),
                                'end_char': int(e_char)
                            }

                            document = Document(
                                client_id=client_id,
                                filename=filename,
                                file_type=file_type,
                                chunk_text=chunk,
                                chunk_index=chunk_idx,
                                embedding=emb_vec,
                                doc_metadata=json.dumps(metadata)
                            )
                            try:
                                db.add(document)
                                db.flush()
                            except Exception as e:
                                logger.warning("VECTOR INSERT FAILED: %s", e)
                            chunk_idx += 1
                            processed_chunks += 1
                        except Exception as e:
                            logger.exception("Error processing chunk %s on page %s: %s", chunk_idx, page_num, e)
                    db.commit()
                    logger.debug("Committed %s chunks for %s (page %s/%s)", processed_chunks, filename, page_num+1, total_pages)
                    try:
                        _append_debug_log(f"DB_COMMIT: filename={filename}, page={page_num+1}, committed_chunks={processed_chunks}")
                    except Exception:
                        pass
            
            doc.close()
            db.commit()
            try:
                _append_debug_log(f"DB_COMMIT: filename={filename}, final_commit_after_pdf_close, committed_chunks={processed_chunks}")
            except Exception:
                pass
            
        elif file_type in ('txt', 'docx'):
            try:
                text = extract_text_from_txt(file_path) if file_type == 'txt' else extract_text_from_docx(file_path)
            except Exception as e:
                _append_debug_log(f"EXTRACT_FAILED: filename={filename}, client_id={client_id}, error={e}")
                raise
            if not text or not text.strip():
                _append_debug_log(f"EXTRACT_EMPTY: filename={filename}, client_id={client_id}")
                raise ValueError("No text content found in file")
            
            text = re.sub(r'\s+', ' ', text).strip()
            chunks = chunk_text(text, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP)
            try:
                _append_debug_log(f"CHUNKS_CREATED: filename={filename}, client_id={client_id}, chunk_count={len(chunks)}")
            except Exception:
                pass
            
            processed_chunks = 0
            BATCH_SIZE = int(os.getenv('EMBED_BATCH_SIZE', '16'))
            batch = []

            for idx, (chunk, start_char, end_char) in enumerate(chunks):
                batch.append((chunk, start_char, end_char, idx))

                if len(batch) >= BATCH_SIZE:
                    texts = [b[0] for b in batch]
                    embeddings = gemini_service.generate_embeddings(texts)
                    for emb_vec, (chunk, s_char, e_char, c_idx) in zip(embeddings, batch):
                        try:
                            if isinstance(emb_vec, dict):
                                emb_vec = emb_vec.get('embedding', emb_vec)
                            try:
                                emb_vec = emb_vec.tolist()
                            except:
                                pass
                            emb_vec = [float(x) for x in emb_vec]

                            metadata = {
                                'page': 1,
                                'start_char': int(s_char),
                                'end_char': int(e_char)
                            }

                            document = Document(
                                client_id=client_id,
                                filename=filename,
                                file_type=file_type,
                                chunk_text=chunk,
                                chunk_index=c_idx,
                                embedding=emb_vec,
                                doc_metadata=json.dumps(metadata)
                            )
                            try:
                                # Debug: log the client_id and chunk index right before DB insert
                                try:
                                    logger.debug("DEBUG INSERT: client_id=%s, filename=%s, chunk_index=%s", client_id, filename, c_idx)
                                    _append_debug_log(f"DEBUG_INSERT: client_id={client_id}, filename={filename}, chunk_index={c_idx}")
                                except Exception:
                                    pass
                                db.add(document)
                                db.flush()
                            except Exception as e:
                                logger.warning("VECTOR INSERT FAILED: %s", e)
                            processed_chunks += 1
                        except Exception as e:
                            logger.exception("Error processing text chunk %s: %s", c_idx, e)

                    db.commit()
                    try:
                        _append_debug_log(f"DB_COMMIT: filename={filename}, committed_chunks={processed_chunks}")
                    except Exception:
                        pass
                    batch = []

            # Flush remaining
            if batch:
                texts = [b[0] for b in batch]
                embeddings = gemini_service.generate_embeddings(texts)
                for emb_vec, (chunk, s_char, e_char, c_idx) in zip(embeddings, batch):
                    try:
                        if isinstance(emb_vec, dict):
                            emb_vec = emb_vec.get('embedding', emb_vec)
                        try:
                            emb_vec = emb_vec.tolist()
                        except:
                            pass
                        emb_vec = [float(x) for x in emb_vec]

                        metadata = {
                            'page': 1,
                            'start_char': int(s_char),
                            'end_char': int(e_char)
                        }

                        document = Document(
                            client_id=client_id,
                            filename=filename,
                            file_type=file_type,
                            chunk_text=chunk,
                            chunk_index=c_idx,
                            embedding=emb_vec,
                            doc_metadata=json.dumps(metadata)
                        )
                        try:
                            # Debug: log the client_id and chunk index right before DB insert
                            try:
                                logger.debug("DEBUG INSERT: client_id=%s, filename=%s, chunk_index=%s", client_id, filename, c_idx)
                                _append_debug_log(f"DEBUG_INSERT: client_id={client_id}, filename={filename}, chunk_index={c_idx}")
                            except Exception:
                                pass
                            db.add(document)
                            db.flush()
                        except Exception as e:
                            logger.warning("VECTOR INSERT FAILED: %s", e)
                        processed_chunks += 1
                    except Exception as e:
                        logger.exception("Error processing text chunk %s: %s", c_idx, e)
                db.commit()
                try:
                    _append_debug_log(f"DB_COMMIT: filename={filename}, final_committed_chunks={processed_chunks}")
                except Exception:
                    pass
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Debug: report completion and association
        try:
            logger.info("process_document COMPLETE: filename=%s, processed_chunks=%s, client_id=%s", filename, processed_chunks, client_id)
        except Exception:
            pass

        # update audit row as completed
        try:
            if req_id:
                audit = db.query(EnqueueAudit).filter(EnqueueAudit.task_id == req_id).first()
            else:
                audit = db.query(EnqueueAudit).filter(EnqueueAudit.file_path == file_path, EnqueueAudit.filename == filename).order_by(EnqueueAudit.created_at.desc()).first()
            if audit:
                audit.status = 'completed'
                audit.completed_at = func.now()
                db.add(audit)
                db.commit()
                _append_debug_log(f"AUDIT_COMPLETED: task_id={audit.task_id}, file={filename}, client_id={audit.client_id}")
        except Exception:
            try:
                db.rollback()
            except Exception:
                pass

        # Clean up uploaded file
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            logger.warning("Error removing file %s: %s", file_path, e)

        return {
            "status": "success",
            "filename": filename,
            "processed_chunks": processed_chunks
        }
        
    except Exception as e:
        try:
            db.rollback()
        except Exception:
            pass
        logger.exception("Error processing document: %s", e)
        try:
            tb = traceback.format_exc()
            _append_debug_log(f"ERROR: task failed filename={filename}, file_path={file_path}, client_id={client_id}, error={e}, traceback={tb}")
        except Exception:
            pass
        
        # Retry logic
        try:
            raise self.retry(exc=e, countdown=60)
        except self.MaxRetriesExceededError:
            return {
                "status": "failed",
                "filename": filename,
                "error": str(e)
            }
    finally:
        db.close()

@celery_app.task
def cleanup_old_documents(days: int = 30) -> dict:
    """
    Cleanup documents older than specified days
    
    Args:
        days: Number of days to keep documents
        
    Returns:
        Dictionary with cleanup results
    """
    db: Session = SessionLocal()
    
    try:
        from datetime import datetime, timedelta
        cutoff_date = datetime.utcnow() - timedelta(days=days)
        
        deleted_count = db.query(Document).filter(
            Document.created_at < cutoff_date
        ).delete()
        
        db.commit()
        
        return {
            "status": "success",
            "deleted_count": deleted_count
        }
    except Exception as e:
        db.rollback()
        return {
            "status": "failed",
            "error": str(e)
        }
    finally:
        db.close()
